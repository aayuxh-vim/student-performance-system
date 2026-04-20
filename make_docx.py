from docx import Document
from docx.shared import Pt, Inches

def create_report():
    doc = Document()
    
    # Title
    doc.add_heading('Mini Project Report: Student Performance System', 0)
    
    # Abstract
    doc.add_heading('Abstract', level=1)
    doc.add_paragraph(
        "This project presents the design and implementation of a simplified, yet highly scalable Student Performance System aimed at modernizing the management of academic data. The system allows administrators to seamlessly handle departments, teachers, students, courses, and course enrollments. By providing a centralized, secure repository, the system ensures accurate grade tracking and calculates student performance metrics efficiently. The backend relies on PostgreSQL (via Supabase) and implements robust constraints, triggers, and PL/pgSQL functions to automate audit logging and grade calculations. The frontend is built using Flask and vanilla HTML/CSS, providing a sleek, dark-mode user interface. The project serves as a proof of concept for a lightweight, easily deployable educational management platform that eliminates data redundancy and enhances administrative workflows."
    )
    
    # Table of Contents
    doc.add_heading('Table of Contents', level=1)
    toc = [
        "1. Introduction",
        "    1.1 Overview",
        "    1.2 Motivation",
        "    1.3 Objectives",
        "2. Problem definition",
        "    2.1 Existing System Drawbacks",
        "    2.2 Proposed System Solutions",
        "3. Tools and Technologies used",
        "    3.1 Frontend Technologies",
        "    3.2 Backend Technologies",
        "    3.3 Database Architecture",
        "4. Database Design (ER diagram)",
        "    4.1 Entity Identification",
        "    4.2 Relationship Mapping",
        "5. Database schema",
        "    5.1 Table Definitions",
        "    5.2 Primary and Foreign Keys",
        "6. Normalization details",
        "    6.1 First Normal Form (1NF)",
        "    6.2 Second Normal Form (2NF)",
        "    6.3 Third Normal Form (3NF)",
        "7. DDL (Data Definition Language)",
        "    7.1 Schema Creation",
        "    7.2 Constraints",
        "8. DML along with the results of the queries",
        "    8.1 Data Insertion",
        "    8.2 Data Retrieval",
        "    8.3 Data Modification",
        "9. DCL (Data Control Language)",
        "    9.1 User Privileges",
        "    9.2 Security Policies",
        "10. Triggers",
        "    10.1 Audit Logging",
        "    10.2 Trigger Implementation",
        "11. PLSQL procedure/function",
        "    11.1 Grade Calculation Logic",
        "12. Frontend GUI screenshots",
        "    12.1 User Authentication Interface",
        "    12.2 Administrator Dashboard",
        "13. Conclusion",
        "    13.1 Project Summary",
        "    13.2 Future Enhancements",
        "14. References in IEEE format"
    ]
    for item in toc:
        doc.add_paragraph(item)
        
    # Abbreviations
    doc.add_heading('List of Abbreviations', level=1)
    abbr = [
        "DDL: Data Definition Language",
        "DML: Data Manipulation Language",
        "DCL: Data Control Language",
        "ER: Entity-Relationship",
        "PL/pgSQL: Procedural Language/PostgreSQL",
        "UI: User Interface",
        "RBAC: Role-Based Access Control",
        "1NF / 2NF / 3NF: First / Second / Third Normal Form"
    ]
    for a in abbr:
        doc.add_paragraph(a, style='List Bullet')
        
    # Figures & Tables
    doc.add_heading('List of Figures & Tables', level=1)
    doc.add_paragraph("Figure 1: Entity-Relationship Diagram", style='List Bullet')
    doc.add_paragraph("Figure 2: Frontend GUI Demo (Screenshots)", style='List Bullet')
    doc.add_paragraph("Table 1: Database Schema Overview", style='List Bullet')
    
    # 1. Introduction
    doc.add_heading('1. Introduction', level=1)
    doc.add_heading('1.1 Overview', level=2)
    doc.add_paragraph("The Student Performance System is a comprehensive web-based application designed to facilitate the administrative tasks of an educational institution. It provides a central hub for managing the academic lifecycle, from creating departments to logging final grades.")
    doc.add_heading('1.2 Motivation', level=2)
    doc.add_paragraph("The motivation behind this project is the growing need for educational institutions to transition from manual, error-prone paper records to digitized, centralized database systems. Many small to medium-scale institutions lack access to affordable, streamlined software, relying instead on cumbersome spreadsheets.")
    doc.add_heading('1.3 Objectives', level=2)
    doc.add_paragraph("The primary objective is to develop a lightweight system capable of securely managing core academic entities. The system emphasizes minimal setup, robust backend data integrity using advanced database features like triggers and PL/pgSQL functions, and an intuitive, modern user interface that requires zero training to operate.")
    
    # 2. Problem Definition
    doc.add_heading('2. Problem definition', level=1)
    doc.add_heading('2.1 Existing System Drawbacks', level=2)
    doc.add_paragraph("Currently, many small-scale academic programs struggle with fragmented data management. Student records, course enrollments, and grading are often siloed across different applications. This fragmentation leads to severe data inconsistencies, difficulty in calculating aggregate metrics like average grades, and an inability to track historical changes.")
    doc.add_heading('2.2 Proposed System Solutions', level=2)
    doc.add_paragraph("The problem requires a cohesive relational database solution integrated with a simple web interface. By enforcing strict foreign key constraints and normalization, the proposed system ensures that data remains consistent. The introduction of automated triggers ensures that every critical action is logged for auditing, solving the issue of untraceable historical changes.")
    
    # 3. Tools
    doc.add_heading('3. Tools and Technologies used', level=1)
    doc.add_heading('3.1 Frontend Technologies', level=2)
    doc.add_paragraph("HTML5: For structuring the web pages. Vanilla CSS3: Used to create a modern, 'glassmorphism' dark-mode aesthetic without the bloat of heavy CSS frameworks.")
    doc.add_heading('3.2 Backend Technologies', level=2)
    doc.add_paragraph("Python (3.x): The primary programming language. Flask: A lightweight WSGI web framework. Psycopg: A highly efficient PostgreSQL adapter.")
    doc.add_heading('3.3 Database Architecture', level=2)
    doc.add_paragraph("PostgreSQL: An advanced, enterprise-class open-source relational database engine hosted via Supabase.")
        
    # 4. ER
    doc.add_heading('4. Database Design (ER diagram)', level=1)
    doc.add_heading('4.1 Entity Identification', level=2)
    doc.add_paragraph("The core entities identified for the system are Department, Teacher, Student, Course, Enrollment, and Grade. An auxiliary entity, Audit Log, is used purely for systemic tracking.")
    doc.add_heading('4.2 Relationship Mapping', level=2)
    doc.add_paragraph("A Department can host multiple Teachers and Students (1:N). A Teacher can teach multiple Courses (1:N). A Student can enroll in multiple Courses, and a Course can have multiple Students (M:N, resolved by the Enrollment table). An Enrollment receives exactly one final Grade (1:1).")
    
    # 5. Schema
    doc.add_heading('5. Database schema', level=1)
    doc.add_heading('5.1 Table Definitions', level=2)
    doc.add_paragraph("Tables include: app_user (Admin Login), department, teacher, student, course, enrollment, grade, and audit_log.")
    doc.add_heading('5.2 Primary and Foreign Keys', level=2)
    doc.add_paragraph("Every table features an auto-incrementing surrogate Primary Key (e.g., student_id). Foreign keys enforce referential integrity with cascading deletes (ON DELETE CASCADE) implemented where appropriate to prevent orphaned records.")
    
    # 6. Normalization
    doc.add_heading('6. Normalization details', level=1)
    doc.add_heading('6.1 First Normal Form (1NF)', level=2)
    doc.add_paragraph("Every table complies with 1NF by ensuring that all attributes are atomic. There are no arrays or comma-separated lists stored within a single column. Every row is uniquely identifiable.")
    doc.add_heading('6.2 Second Normal Form (2NF)', level=2)
    doc.add_paragraph("The schema satisfies 1NF, and all non-key attributes are fully functionally dependent on the entire primary key. Because every table uses a single-column surrogate primary key, there are inherently no partial dependencies.")
    doc.add_heading('6.3 Third Normal Form (3NF)', level=2)
    doc.add_paragraph("The schema satisfies 2NF, and there are no transitive dependencies. Non-key attributes depend strictly on the primary key and nothing else. A department_id foreign key is used to reference the department table rather than storing department strings.")
    
    # 7. DDL
    doc.add_heading('7. DDL (Data Definition Language)', level=1)
    doc.add_heading('7.1 Schema Creation', level=2)
    doc.add_paragraph("CREATE TABLE student ( student_id SERIAL PRIMARY KEY, name TEXT NOT NULL, email TEXT, department_id INT REFERENCES department(department_id) ON DELETE SET NULL );\nCREATE TABLE grade ( grade_id SERIAL PRIMARY KEY, enrollment_id INT NOT NULL REFERENCES enrollment(enrollment_id) ON DELETE CASCADE, score NUMERIC(5, 2) NOT NULL );")
    doc.add_heading('7.2 Constraints', level=2)
    doc.add_paragraph("The UNIQUE(student_id, course_id) constraint ensures a student cannot be enrolled in the same course twice. NOT NULL constraints prevent incomplete entries.")

    # 8. DML
    doc.add_heading('8. DML (Data Manipulation Language)', level=1)
    doc.add_heading('8.1 Data Insertion', level=2)
    doc.add_paragraph("INSERT INTO department (name) VALUES ('Computer Science'); -> Result: department_id = 1")
    doc.add_heading('8.2 Data Retrieval', level=2)
    doc.add_paragraph("SELECT s.name, c.name as course, g.score FROM student s JOIN enrollment e ON s.student_id = e.student_id JOIN course c ON e.course_id = c.course_id JOIN grade g ON e.enrollment_id = g.enrollment_id; -> Result: Grace Hopper | Algorithms 101 | 98.50")
    doc.add_heading('8.3 Data Modification', level=2)
    doc.add_paragraph("UPDATE student SET email = 'grace.h@cs.edu' WHERE student_id = 1; -> Result: 1 row updated")
    
    # 9. DCL
    doc.add_heading('9. DCL (Data Control Language)', level=1)
    doc.add_heading('9.1 User Privileges', level=2)
    doc.add_paragraph("GRANT ALL ON SCHEMA public TO postgres; GRANT SELECT, INSERT, UPDATE, DELETE ON ALL TABLES IN SCHEMA public TO public;")
    doc.add_heading('9.2 Security Policies', level=2)
    doc.add_paragraph("By explicitly defining GRANT permissions, we ensure that the application layer only has the necessary access to manipulate data (CRUD operations), preventing unauthorized structural changes.")
    
    # 10. Triggers
    doc.add_heading('10. Triggers', level=1)
    doc.add_heading('10.1 Audit Logging', level=2)
    doc.add_paragraph("An audit logging trigger was implemented to automatically record whenever a new grade is inserted into the system. This provides a tamper-proof historical log.")
    doc.add_heading('10.2 Trigger Implementation', level=2)
    doc.add_paragraph("CREATE TRIGGER after_grade_insert AFTER INSERT ON grade FOR EACH ROW EXECUTE FUNCTION log_grade_insert();")
    
    # 11. PLSQL
    doc.add_heading('11. PLSQL procedure/function', level=1)
    doc.add_heading('11.1 Grade Calculation Logic', level=2)
    doc.add_paragraph("A procedural function was designed to dynamically calculate the average score of a student across all enrolled courses directly at the database level. \nCREATE OR REPLACE FUNCTION get_student_average(p_student_id INT) RETURNS NUMERIC AS $$ ... $$ LANGUAGE plpgsql;")
    
    # 12. Screenshots
    doc.add_heading('12. Frontend GUI screenshots', level=1)
    doc.add_heading('12.1 User Authentication Interface', level=2)
    doc.add_paragraph("The login interface ensures that only authorized administrators can access the system.")
    doc.add_heading('12.2 Administrator Dashboard', level=2)
    doc.add_paragraph("The dashboard provides a high-level statistical overview of the database and displays a comprehensive roster.")
    
    # 13. Conclusion
    doc.add_heading('13. Conclusion', level=1)
    doc.add_heading('13.1 Project Summary', level=2)
    doc.add_paragraph("The implementation successfully demonstrates how relational databases can be paired with lightweight web frameworks to solve data fragmentation. The inclusion of PostgreSQL Triggers ensures an automated audit trail for grading, while PL/pgSQL functions optimize aggregate data calculation.")
    doc.add_heading('13.2 Future Enhancements', level=2)
    doc.add_paragraph("Future iterations could implement dynamic Role-Based Access Control (RBAC) allowing individual teachers to log in and grade only their assigned courses, and data export functionality.")
    
    # 14. References
    doc.add_heading('14. References in IEEE format', level=1)
    doc.add_paragraph("[1] PostgreSQL Global Development Group, 'PostgreSQL 15.0 Documentation,' PostgreSQL, 2022.")
    doc.add_paragraph("[2] Pallets Projects, 'Flask Documentation (3.0.x),' Flask, 2023.")
    
    doc.save('Project_Report.docx')

if __name__ == '__main__':
    create_report()
